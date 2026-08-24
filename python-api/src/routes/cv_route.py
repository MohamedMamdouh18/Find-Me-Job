import os

from fastapi import APIRouter, Depends, File, HTTPException, UploadFile
from fastapi.responses import FileResponse
from pydantic import BaseModel
from docx import Document
from sqlmodel import Session

from ..shared import CV_PATH
from ..database import get_session
from ..database.repositories import CVKeywordsRepository

cv_router = APIRouter(prefix="/api/cv", tags=["cv"])


class KeywordsRequest(BaseModel):
    cv_hash: str
    keywords: str


def _docx_text(doc) -> str:
    """python-docx keeps table cells out of `doc.paragraphs`, and CVs routinely put
    skills and dates in tables, so reading paragraphs alone silently drops them."""
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n".join(parts)


@cv_router.get("")
def get_cv():
    try:
        doc = Document(CV_PATH)
        text = _docx_text(doc)
        return {"cv_text": text}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Could not read CV: {str(e)}")


MAX_CV_BYTES = 10 * 1024 * 1024
DOCX_MAGIC = b"PK\x03\x04"  # .docx is a zip container


@cv_router.post("/upload")
async def upload_cv(file: UploadFile = File(...)):
    """Replace cv.docx from the dashboard so users never have to touch the filesystem."""
    if not (file.filename or "").lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="CV must be a .docx file")

    content = await file.read()
    if len(content) > MAX_CV_BYTES:
        raise HTTPException(status_code=413, detail="CV is larger than 10 MB")
    if not content.startswith(DOCX_MAGIC):
        raise HTTPException(status_code=400, detail="File is not a valid .docx document")

    # Parse before overwriting so a corrupt upload cannot destroy a working CV.
    try:
        import io

        doc = Document(io.BytesIO(content))
        text = _docx_text(doc)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not read .docx: {e}")

    if not text.strip():
        raise HTTPException(status_code=400, detail="CV contains no readable text")

    try:
        # Written in place: CV_PATH is a single-file bind mount, so replacing the
        # inode via rename would break the mount.
        with open(CV_PATH, "wb") as f:
            f.write(content)
    except OSError as e:
        raise HTTPException(status_code=500, detail=f"Could not save CV: {e}")

    return {"status": "ok", "bytes": len(content), "characters": len(text)}


@cv_router.get("/info")
def cv_info():
    """Size and mtime of the current CV, for the settings screen."""
    if not os.path.isfile(CV_PATH):
        return {"exists": False}
    stat = os.stat(CV_PATH)
    return {"exists": True, "bytes": stat.st_size, "modified_at": stat.st_mtime}


@cv_router.get("/file")
def download_cv():
    """Hand back the actual .docx so the dashboard can offer 'Download current CV'."""
    if not os.path.isfile(CV_PATH):
        raise HTTPException(status_code=404, detail="No CV uploaded")
    return FileResponse(
        CV_PATH,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="cv.docx",
    )


@cv_router.get("/check/{cv_hash}")
def check_cv_hash(cv_hash: str, session: Session = Depends(get_session)):
    repo = CVKeywordsRepository(session)
    return {"exists": repo.hash_exists(cv_hash)}


@cv_router.get("/keywords")
def get_keywords(session: Session = Depends(get_session)):
    repo = CVKeywordsRepository(session)
    row = repo.get_latest()
    if not row:
        return {"keywords": None, "cv_hash": None, "updated_at": None}
    return {"keywords": row.keywords, "cv_hash": row.cv_hash, "updated_at": row.updated_at}


@cv_router.post("/keywords")
def save_keywords(body: KeywordsRequest, session: Session = Depends(get_session)):
    repo = CVKeywordsRepository(session)
    repo.save(body.cv_hash, body.keywords)
    session.commit()
    return {"status": "ok"}
