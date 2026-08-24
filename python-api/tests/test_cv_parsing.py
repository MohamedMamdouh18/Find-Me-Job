import docx
from src.routes.cv_route import _docx_text


def test_paragraphs_only():
    doc = docx.Document()
    doc.add_paragraph("Software Engineer")
    doc.add_paragraph("Python, FastAPI, SQL")
    assert _docx_text(doc) == "Software Engineer\nPython, FastAPI, SQL"


def test_tables_only():
    doc = docx.Document()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Skill: Python"
    table.cell(0, 1).text = "Experience: 5 years"
    assert _docx_text(doc) == "Skill: Python\nExperience: 5 years"


def test_paragraphs_and_tables():
    doc = docx.Document()
    doc.add_paragraph("Candidate Name")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Backend Developer"
    table.cell(0, 1).text = "Full Time"
    doc.add_paragraph("Education: BS CS")
    assert (
        _docx_text(doc)
        == "Candidate Name\nEducation: BS CS\nBackend Developer\nFull Time"
    )
